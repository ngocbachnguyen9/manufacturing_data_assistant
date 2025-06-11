# In src/tools/packing_list_parser_tool.py

import docx # Add this import
import os
from .base_tool import BaseTool
from typing import Dict, Any

class PackingListParserTool(BaseTool):
    """Tool for parsing .docx Packing Lists to find the associated Order ID."""

    def __init__(self, datasets: Dict[str, Any]):
        super().__init__(datasets)
        self.doc_path = "data/generated_documents/packing_lists"

    def run(self, packing_list_id: str, **kwargs) -> Dict[str, str]:
        """
        Finds a packing list by its ID and extracts the OrderNumber.
        """
        doc_filename = f"PackingList-{packing_list_id}.docx" # UPDATED extension
        doc_full_path = os.path.join(self.doc_path, doc_filename)

        if not os.path.exists(doc_full_path):
            return {"error": f"Packing List not found: {doc_filename}"}

        try:
            document = docx.Document(doc_full_path)
            # This is a simple parsing strategy. It assumes the text
            # 'Order Number:' appears in the document.
            for para in document.paragraphs:
                if "Order Number:" in para.text:
                    # Assumes format is "Order Number: ORBOX0011"
                    parts = para.text.split(":")
                    if len(parts) > 1:
                        order_id = parts[1].strip()
                        return {"order_id": order_id, "source_document": doc_filename}

            # Fallback if not found in paragraphs (e.g., it's in a table)
            for table in document.tables:
                for row in table.rows:
                    for i, cell in enumerate(row.cells):
                        if "Order Number" in cell.text and i + 1 < len(row.cells):
                            order_id = row.cells[i + 1].text.strip()
                            return {"order_id": order_id, "source_document": doc_filename}

            return {"error": f"Could not find 'Order Number' in {doc_filename}."}
        except Exception as e:
            return {"error": f"Failed to parse DOCX {doc_filename}: {e}"}